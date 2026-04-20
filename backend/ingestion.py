"""
ingestion.py — Full document ingestion pipeline.

Pipeline stages (in order):
  1. parse_file()     — extract raw text + page metadata from PDF/DOCX/CSV/TXT
  2. chunk_text()     — split text into 300-word chunks with 50-word overlap
  3. embed_chunks()   — call Ollama nomic-embed-text to get 768-dim vectors
  4. store_document() — persist Document → Chunks → Embeddings in one DB transaction
  5. build_bm25_index() — extract BM25 term frequencies and store in chunk metadata

Entry point: ingest_file(file_path, file_name, db_session)
"""

from __future__ import annotations  # enables X | Y union syntax on Python 3.9

import csv
import io
import json
import math
import re
import tempfile
import time
from collections import Counter
from pathlib import Path
from typing import Any, Optional

import httpx
import fitz                          # PyMuPDF
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from database import Chunk, Document, Embedding

# ---------------------------------------------------------------------------
# Configuration constants
# ---------------------------------------------------------------------------

CHUNK_SIZE_WORDS    = 300   # target words per chunk
CHUNK_OVERLAP_WORDS = 50    # words shared between consecutive chunks
OLLAMA_BASE_URL     = "http://localhost:11434"
EMBED_MODEL         = "nomic-embed-text"
OLLAMA_TIMEOUT      = 60    # seconds — embedding a long chunk can take a moment

# ---------------------------------------------------------------------------
# Stage 1 — File parsing
# ---------------------------------------------------------------------------

def parse_pdf(file_path: str) -> list[dict[str, Any]]:
    """
    Parse a PDF file using PyMuPDF.
    Returns a list of page dicts: {"text": str, "page_number": int}
    """
    pages = []
    try:
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:  # skip blank pages
                pages.append({"text": text, "page_number": page_num})
        doc.close()
        print(f"  [parse] PDF: extracted {len(pages)} non-empty pages")
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF '{file_path}': {e}") from e
    return pages


def parse_docx(file_path: str) -> list[dict[str, Any]]:
    """
    Parse a DOCX file using python-docx.
    Groups paragraphs into pseudo-pages of ~500 words each.
    Returns a list of page dicts: {"text": str, "page_number": int}
    """
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX '{file_path}': {e}") from e

    # Group paragraphs into pseudo-pages (~500 words each)
    pages = []
    current_words: list[str] = []
    page_num = 1
    for para in paragraphs:
        words = para.split()
        current_words.extend(words)
        if len(current_words) >= 500:
            pages.append({"text": " ".join(current_words), "page_number": page_num})
            current_words = []
            page_num += 1
    if current_words:
        pages.append({"text": " ".join(current_words), "page_number": page_num})

    print(f"  [parse] DOCX: extracted {len(paragraphs)} paragraphs → {len(pages)} pseudo-pages")
    return pages


def parse_csv(file_path: str) -> list[dict[str, Any]]:
    """
    Parse a CSV file using the standard csv module.
    Each row is converted to a natural-language sentence: "column: value, column: value, ..."
    Rows are grouped into chunks of 50 rows per pseudo-page.
    Returns a list of page dicts: {"text": str, "page_number": int}
    """
    ROWS_PER_PAGE = 50
    try:
        with open(file_path, newline="", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            rows = list(reader)
            headers = reader.fieldnames or []
    except Exception as e:
        raise RuntimeError(f"Failed to parse CSV '{file_path}': {e}") from e

    pages = []
    for page_idx, start in enumerate(range(0, len(rows), ROWS_PER_PAGE), start=1):
        batch = rows[start : start + ROWS_PER_PAGE]
        lines = []
        for row in batch:
            # Build a readable sentence from each row
            parts = [f"{k}: {v}" for k, v in row.items() if v and v.strip()]
            lines.append(", ".join(parts))
        text = "\n".join(lines)
        pages.append({
            "text": text,
            "page_number": page_idx,
            "metadata": {
                "row_start": start + 1,
                "row_end": start + len(batch),
                "headers": headers,
            },
        })

    print(f"  [parse] CSV: {len(rows)} rows → {len(pages)} pseudo-pages")
    return pages


def parse_txt(file_path: str) -> list[dict[str, Any]]:
    """
    Parse a plain-text file.
    Splits into pseudo-pages of ~500 words each.
    Returns a list of page dicts: {"text": str, "page_number": int}
    """
    try:
        with open(file_path, encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        raise RuntimeError(f"Failed to parse TXT '{file_path}': {e}") from e

    words = content.split()
    pages = []
    for page_idx, start in enumerate(range(0, len(words), 500), start=1):
        text = " ".join(words[start : start + 500])
        pages.append({"text": text, "page_number": page_idx})

    print(f"  [parse] TXT: {len(words)} words → {len(pages)} pseudo-pages")
    return pages


def parse_file(file_path: str, file_type: str) -> list[dict[str, Any]]:
    """
    Dispatch to the correct parser based on file_type.
    Returns a list of page dicts with at least {"text": str, "page_number": int}.
    """
    parsers = {
        "pdf":  parse_pdf,
        "docx": parse_docx,
        "csv":  parse_csv,
        "txt":  parse_txt,
    }
    if file_type not in parsers:
        raise ValueError(f"Unsupported file type: '{file_type}'. Supported: {list(parsers.keys())}")
    return parsers[file_type](file_path)


# ---------------------------------------------------------------------------
# Stage 2 — Chunking
# ---------------------------------------------------------------------------

def chunk_text(
    text: str,
    page_number: int,
    extra_metadata: Optional[dict] = None,
    chunk_size: int = CHUNK_SIZE_WORDS,
    overlap: int = CHUNK_OVERLAP_WORDS,
) -> list[dict[str, Any]]:
    """
    Split text into overlapping word-based chunks.

    Args:
        text:           The raw text to chunk.
        page_number:    Source page number (stored in metadata).
        extra_metadata: Any additional metadata to merge in (e.g. CSV headers).
        chunk_size:     Target number of words per chunk (default 300).
        overlap:        Number of words to repeat at the start of the next chunk (default 50).

    Returns:
        List of chunk dicts: {"content": str, "page_number": int, "chunk_index": int, "metadata": dict}
    """
    words = text.split()
    if not words:
        return []

    chunks = []
    start = 0
    chunk_index = 0

    while start < len(words):
        end = start + chunk_size
        chunk_words = words[start:end]
        content = " ".join(chunk_words)

        metadata: dict[str, Any] = {"page_number": page_number, "word_count": len(chunk_words)}
        if extra_metadata:
            metadata.update(extra_metadata)

        chunks.append({
            "content":     content,
            "page_number": page_number,
            "chunk_index": chunk_index,
            "metadata":    metadata,
        })

        chunk_index += 1
        # Advance by (chunk_size - overlap) so the next chunk shares `overlap` words
        step = chunk_size - overlap
        start += step

    return chunks


def chunk_pages(pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """
    Apply chunk_text() to every page and return a flat list of all chunks,
    with globally sequential chunk_index values.
    """
    all_chunks: list[dict[str, Any]] = []
    global_index = 0

    for page in pages:
        page_chunks = chunk_text(
            text=page["text"],
            page_number=page.get("page_number", 0),
            extra_metadata=page.get("metadata"),
        )
        for chunk in page_chunks:
            chunk["chunk_index"] = global_index
            all_chunks.append(chunk)
            global_index += 1

    return all_chunks


# ---------------------------------------------------------------------------
# Stage 3 — Embedding via Ollama
# ---------------------------------------------------------------------------

def embed_text(text: str) -> list[float]:
    """
    Call the Ollama /api/embeddings endpoint to get a 768-dim vector for `text`.
    Uses nomic-embed-text model.

    Raises RuntimeError if Ollama is unreachable or returns an error.
    """
    url = f"{OLLAMA_BASE_URL}/api/embeddings"
    payload = {"model": EMBED_MODEL, "prompt": text}

    try:
        response = httpx.post(url, json=payload, timeout=OLLAMA_TIMEOUT)
        response.raise_for_status()
    except httpx.ConnectError:
        raise RuntimeError(
            f"Cannot connect to Ollama at {OLLAMA_BASE_URL}. "
            "Make sure Ollama is running: `ollama serve`"
        )
    except httpx.HTTPStatusError as e:
        raise RuntimeError(f"Ollama returned HTTP {e.response.status_code}: {e.response.text}") from e

    data = response.json()
    embedding = data.get("embedding")
    if not embedding:
        raise RuntimeError(f"Ollama response missing 'embedding' field: {data}")

    return embedding


def embed_chunks(chunks: list[dict[str, Any]]) -> list[list[float]]:
    """
    Embed every chunk in sequence, printing progress every 10 chunks.
    Returns a list of 768-dim float vectors, one per chunk.
    """
    vectors = []
    total = len(chunks)

    for i, chunk in enumerate(chunks):
        vector = embed_text(chunk["content"])
        vectors.append(vector)

        # Progress logging every 10 chunks (and on the last one)
        if (i + 1) % 10 == 0 or (i + 1) == total:
            print(f"  [embed] {i + 1}/{total} chunks embedded")

    return vectors


# ---------------------------------------------------------------------------
# Stage 4 — BM25 term index
# ---------------------------------------------------------------------------

def tokenize(text: str) -> list[str]:
    """
    Simple whitespace + punctuation tokenizer.
    Lowercases and strips non-alphanumeric characters.
    """
    return re.findall(r"[a-z0-9]+", text.lower())


def compute_bm25_terms(chunks: list[dict[str, Any]]) -> list[dict[str, float]]:
    """
    Compute BM25 term frequency × IDF scores for each chunk.

    BM25 formula used:
        score(t, d) = IDF(t) × (tf × (k1 + 1)) / (tf + k1 × (1 - b + b × |d| / avgdl))

    Returns a list of dicts mapping term → BM25 score, one dict per chunk.
    These are stored in chunk metadata under the key "bm25_terms".
    """
    k1 = 1.5
    b  = 0.75

    # Tokenize all chunks
    tokenized = [tokenize(c["content"]) for c in chunks]
    doc_lengths = [len(t) for t in tokenized]
    avgdl = sum(doc_lengths) / max(len(doc_lengths), 1)
    N = len(chunks)

    # Document frequency: how many chunks contain each term
    df: Counter = Counter()
    for tokens in tokenized:
        for term in set(tokens):
            df[term] += 1

    bm25_scores_per_chunk = []
    for tokens, dl in zip(tokenized, doc_lengths):
        tf_map = Counter(tokens)
        scores: dict[str, float] = {}
        for term, tf in tf_map.items():
            idf = math.log((N - df[term] + 0.5) / (df[term] + 0.5) + 1)
            numerator   = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / max(avgdl, 1))
            scores[term] = round(idf * numerator / denominator, 6)
        bm25_scores_per_chunk.append(scores)

    return bm25_scores_per_chunk


# ---------------------------------------------------------------------------
# Stage 5 — Database persistence
# ---------------------------------------------------------------------------

def store_document(
    db: Session,
    file_name: str,
    file_type: str,
    chunks: list[dict[str, Any]],
    vectors: list[list[float]],
    bm25_scores: list[dict[str, float]],
) -> Document:
    """
    Persist the full ingestion result in a single database transaction:
      1. Insert Document row
      2. Insert all Chunk rows (with BM25 terms merged into metadata)
      3. Insert all Embedding rows

    Rolls back the entire transaction on any error.
    Returns the saved Document ORM object.
    """
    try:
        # --- Document ---
        doc = Document(name=file_name, type=file_type)
        db.add(doc)
        db.flush()  # get doc.id without committing
        print(f"  [store] Document id={doc.id} created")

        # --- Chunks + Embeddings ---
        for i, (chunk_data, vector, bm25) in enumerate(zip(chunks, vectors, bm25_scores)):
            # Merge BM25 terms into chunk metadata
            meta = chunk_data.get("metadata") or {}
            meta["bm25_terms"] = bm25

            chunk = Chunk(
                document_id    = doc.id,
                content        = chunk_data["content"],
                page_number    = chunk_data.get("page_number"),
                chunk_index    = chunk_data["chunk_index"],
                chunk_metadata = meta,
            )
            db.add(chunk)
            db.flush()  # get chunk.id

            embedding = Embedding(chunk_id=chunk.id, vector=vector)
            db.add(embedding)

            if (i + 1) % 20 == 0 or (i + 1) == len(chunks):
                print(f"  [store] {i + 1}/{len(chunks)} chunks + embeddings staged")

        db.commit()
        db.refresh(doc)
        print(f"  [store] Transaction committed — {len(chunks)} chunks stored")
        return doc

    except Exception as e:
        db.rollback()
        raise RuntimeError(f"Database transaction failed and was rolled back: {e}") from e


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def ingest_file(file_path: str, file_name: str, db: Session) -> dict[str, Any]:
    """
    Run the full ingestion pipeline for a single file.

    Steps:
      1. Detect file type from extension
      2. Parse file into pages
      3. Chunk pages into overlapping text windows
      4. Embed each chunk via Ollama
      5. Compute BM25 term scores
      6. Store everything in Supabase

    Args:
        file_path: Absolute path to the temporary file on disk.
        file_name: Original filename (used for display and type detection).
        db:        SQLAlchemy session (injected by FastAPI dependency).

    Returns:
        {"document_id": int, "chunk_count": int, "file_name": str}
    """
    start_time = time.time()
    ext = Path(file_name).suffix.lstrip(".").lower()

    # Normalize common extensions
    ext_map = {"docx": "docx", "pdf": "pdf", "csv": "csv", "txt": "txt"}
    file_type = ext_map.get(ext)
    if not file_type:
        raise ValueError(f"Unsupported file extension '.{ext}'. Supported: pdf, docx, csv, txt")

    print(f"\n{'='*60}")
    print(f"[ingest] Starting ingestion: {file_name} (type={file_type})")
    print(f"{'='*60}")

    # Stage 1 — Parse
    print("[ingest] Stage 1/5 — Parsing file...")
    pages = parse_file(file_path, file_type)
    if not pages:
        raise ValueError(f"No text content could be extracted from '{file_name}'")
    print(f"  → {len(pages)} pages extracted")

    # Stage 2 — Chunk
    print("[ingest] Stage 2/5 — Chunking text...")
    chunks = chunk_pages(pages)
    if not chunks:
        raise ValueError(f"No chunks produced from '{file_name}' — file may be empty")
    print(f"  → {len(chunks)} chunks created")

    # Stage 3 — Embed
    print("[ingest] Stage 3/5 — Embedding chunks via Ollama...")
    vectors = embed_chunks(chunks)
    print(f"  → {len(vectors)} embeddings generated")

    # Stage 4 — BM25
    print("[ingest] Stage 4/5 — Computing BM25 term scores...")
    bm25_scores = compute_bm25_terms(chunks)
    print(f"  → BM25 index built for {len(bm25_scores)} chunks")

    # Stage 5 — Store
    print("[ingest] Stage 5/5 — Storing in Supabase...")
    doc = store_document(db, file_name, file_type, chunks, vectors, bm25_scores)

    elapsed = round(time.time() - start_time, 2)
    print(f"\n[ingest] ✓ Done in {elapsed}s — document_id={doc.id}, chunks={len(chunks)}")
    print(f"{'='*60}\n")

    return {
        "document_id": doc.id,
        "chunk_count": len(chunks),
        "file_name":   file_name,
    }
